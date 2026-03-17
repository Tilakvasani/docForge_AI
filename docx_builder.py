"""
DocForge AI — docx_builder.py  v5.0  (pure Python, no Node.js)

Features:
  ✅ Professional header:  Company  |  Doc Type  (right-aligned, grey)
  ✅ Footer:               Page N of M  (centred, grey)
  ✅ Title block:          Doc type (large, accent blue) + subtitle line
  ✅ Section headings:     accent blue, bold, underline divider
  ✅ Pipe-format tables  → real Word tables (header shaded, zebra rows)
  ✅ Mermaid flowcharts  → real PNG image embedded in Word (via flowchart_renderer.py)
  ✅ Bullet lines        → indented bullet paragraphs
  ✅ Numbered lines      → indented numbered paragraphs
  ✅ Plain paragraphs
  ✅ Returns raw bytes   — no temp files, no subprocess

Dependencies:
    pip install python-docx matplotlib pillow

Both files must be in the same directory:
    docx_builder.py
    flowchart_renderer.py
"""

import re
import io
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from flowchart_renderer import mermaid_to_png_bytes
    FLOWCHART_RENDERER_AVAILABLE = True
except ImportError:
    FLOWCHART_RENDERER_AVAILABLE = False


# ── Colour palette ─────────────────────────────────────────────────────────────
ACCENT  = RGBColor(0x2E, 0x40, 0x57)
GRAY    = RGBColor(0x88, 0x88, 0x88)
LGRAY   = RGBColor(0xF3, 0xF4, 0xF6)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DARK    = RGBColor(0x22, 0x22, 0x22)
BORDER  = RGBColor(0xCC, 0xCC, 0xCC)
FLOW_BG = RGBColor(0xEE, 0xF2, 0xFF)
FLOW_BD = RGBColor(0x66, 0x7E, 0xEA)


# ─────────────────────────────────────────────────────────────────────────────
#  LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _set_run_font(run, size_pt: float, bold=False, italic=False,
                  color: RGBColor = None, font_name="Arial"):
    run.font.name   = font_name
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _cell_shade(cell, fill_color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _hex(fill_color))
    tcPr.append(shd)


def _cell_borders(cell, color: RGBColor = None):
    bc   = color or BORDER
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _hex(bc))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_page_number_field(paragraph):
    def _field(fld_char_type):
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), fld_char_type)
        return fldChar

    def _instr(text):
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = text
        return instrText

    p = paragraph._p
    r1 = OxmlElement("w:r")
    r1.append(OxmlElement("w:rPr"))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = "Page "
    r1.append(t)
    p.append(r1)

    for part in ("begin", "PAGE", "end"):
        r = OxmlElement("w:r")
        r.append(_field(part) if part in ("begin", "end") else _instr(f" {part} "))
        p.append(r)

    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.set(qn("xml:space"), "preserve")
    t2.text = " of "
    r2.append(t2)
    p.append(r2)

    for part in ("begin", "NUMPAGES", "end"):
        r = OxmlElement("w:r")
        r.append(_field(part) if part in ("begin", "end") else _instr(f" {part} "))
        p.append(r)


def _add_bottom_border(paragraph):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _hex(ACCENT))
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
#  PIPE TABLE → WORD TABLE
# ─────────────────────────────────────────────────────────────────────────────

def _is_table_line(line: str) -> bool:
    return "|" in line

def _is_separator(line: str) -> bool:
    return bool(re.match(r"^\s*[\|\-\s:]+$", line)) and "-" in line

def _parse_pipe_table(lines: List[str]) -> List[List[str]]:
    rows = []
    for line in lines:
        if _is_separator(line):
            continue
        if not _is_table_line(line):
            continue
        parts = line.split("|")
        cells = [c.strip() for c in parts]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows


def _build_word_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table     = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_width = 9360 // col_count

    for ri, row_data in enumerate(rows):
        is_header = (ri == 0)
        tr        = table.rows[ri]
        for ci in range(col_count):
            cell      = tr.cells[ci]
            cell_text = row_data[ci] if ci < len(row_data) else ""
            if is_header:
                _cell_shade(cell, ACCENT)
            elif ri % 2 == 0:
                _cell_shade(cell, LGRAY)
            _cell_borders(cell)
            _cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Twips(col_width)
            p   = cell.paragraphs[0]
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(cell_text)
            _set_run_font(run, size_pt=10, bold=is_header,
                          color=WHITE if is_header else DARK)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


# ─────────────────────────────────────────────────────────────────────────────
#  FLOWCHART → REAL PNG IMAGE
# ─────────────────────────────────────────────────────────────────────────────

def _build_flowchart_image(doc: Document, mermaid_text: str, section_name: str = ""):
    """
    Render Mermaid flowchart as a real PNG image and embed in the Word doc.
    Falls back to a styled text box if flowchart_renderer.py is not installed.
    """
    if FLOWCHART_RENDERER_AVAILABLE:
        try:
            png_bytes    = mermaid_to_png_bytes(mermaid_text, title=section_name, dpi=180)
            image_stream = io.BytesIO(png_bytes)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            p.add_run().add_picture(image_stream, width=Inches(6.2))

            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            _set_run_font(
                cap.add_run(
                    f"Figure: {section_name} — Process Flow Diagram"
                    if section_name else "Figure: Process Flow Diagram"
                ),
                size_pt=8, italic=True, color=GRAY
            )
            return
        except Exception:
            pass

    # Fallback
    _build_flowchart_text_fallback(doc, mermaid_text, section_name)


def _extract_mermaid_steps(mermaid_text: str) -> List[str]:
    steps, seen = [], set()
    pattern = re.compile(r'\w+[\[\({]+([^\]\)\}]+)[\]\)}]+')
    for line in mermaid_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("flowchart") or line.startswith("graph"):
            continue
        for match in pattern.finditer(line):
            label = match.group(1).strip()
            if label and label not in seen and len(label) > 2:
                seen.add(label)
                steps.append(label)
    return steps


def _build_flowchart_text_fallback(doc: Document, mermaid_text: str, section_name: str = ""):
    steps = _extract_mermaid_steps(mermaid_text)
    box   = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell  = box.cell(0, 0)
    _cell_shade(cell, FLOW_BG)
    _cell_borders(cell, FLOW_BD)
    _cell_margins(cell, top=160, bottom=160, left=200, right=200)

    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(6)
    _set_run_font(title_p.add_run("Process Flow Diagram"),
                  size_pt=11, bold=True, color=ACCENT)

    sub_p = cell.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(10)
    _set_run_font(sub_p.add_run(
        f"Process flow for {section_name}. "
        f"View the interactive diagram in DocForge AI."
        if section_name else "View the interactive diagram in DocForge AI."
    ), size_pt=9, italic=True, color=GRAY)

    if steps:
        for idx, step in enumerate(steps):
            if idx > 0:
                arrow_p = cell.add_paragraph()
                arrow_p.paragraph_format.space_after  = Pt(1)
                arrow_p.paragraph_format.left_indent  = Inches(0.2)
                _set_run_font(arrow_p.add_run("↓"), size_pt=10, color=FLOW_BD)
            step_p = cell.add_paragraph()
            step_p.paragraph_format.space_after = Pt(2)
            step_p.paragraph_format.left_indent = Inches(0.2)
            _set_run_font(step_p.add_run(f"  {idx+1}.  "), size_pt=10, bold=True, color=FLOW_BD)
            _set_run_font(step_p.add_run(step), size_pt=10, color=DARK)
    else:
        _set_run_font(cell.add_paragraph().add_run("Steps defined in the flowchart."),
                      size_pt=10, italic=True, color=GRAY)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────────────────────────────────────
#  SECTION CONTENT RENDERER
# ─────────────────────────────────────────────────────────────────────────────

def _render_section_content(doc: Document, content: str, section_name: str = ""):
    """
    Parse section content and add the correct Word elements.

    Priority order:
      1. ```mermaid...```  → real PNG flowchart image
      2. Pipe tables       → Word table
      3. Numbered list     → indented numbered paragraph
      4. Bullet list       → indented bullet paragraph
      5. Empty line        → spacer
      6. Plain text        → paragraph
    """
    # Safety net: auto-wrap bare flowchart blocks if LLM dropped the fences
    if re.search(r'flowchart\s+(?:TD|LR|BT|RL)', content) and "```mermaid" not in content:
        content = re.sub(
            r'(flowchart\s+(?:TD|LR|BT|RL).*?)(\n\n|\Z)',
            lambda m: "```mermaid\n" + m.group(1).rstrip() + "\n```" + m.group(2),
            content,
            flags=re.DOTALL
        )

    mermaid_pattern = re.compile(r"```mermaid(.*?)```", re.DOTALL)
    parts = mermaid_pattern.split(content)

    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            _build_flowchart_image(doc, part.strip(), section_name)
        else:
            _render_plain_content(doc, part)


def _render_plain_content(doc: Document, content: str):
    lines = content.split("\n")
    i     = 0

    while i < len(lines):
        line = lines[i]

        # Pipe table block
        if _is_table_line(line):
            table_lines = []
            while i < len(lines) and (
                _is_table_line(lines[i]) or _is_separator(lines[i])
            ):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_pipe_table(table_lines)
            if rows:
                _build_word_table(doc, rows)
            continue

        stripped = line.strip()

        # Skip stray backtick fences
        if stripped.startswith("```"):
            i += 1
            continue

        # Empty line → spacer
        if not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Numbered list: "1. text" or "1) text"
        num_match = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if num_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            _set_run_font(p.add_run(f"{num_match.group(1)}. {num_match.group(2)}"),
                          size_pt=11, color=DARK)
            i += 1
            continue

        # Bullet: "- text" or "• text"
        bullet_match = re.match(r"^[-•]\s+(.+)$", stripped)
        if bullet_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_after       = Pt(4)
            _set_run_font(p.add_run("•  "), size_pt=11, color=ACCENT)
            _set_run_font(p.add_run(bullet_match.group(1)), size_pt=11, color=DARK)
            i += 1
            continue

        # Clean leftover markdown symbols
        cleaned = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", stripped)
        cleaned = re.sub(r"#{1,6}\s*", "", cleaned).strip()

        if not cleaned:
            i += 1
            continue

        # Plain paragraph
        p   = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _set_run_font(p.add_run(cleaned), size_pt=11, color=DARK)
        i += 1


# ─────────────────────────────────────────────────────────────────────────────
#  HEADER & FOOTER
# ─────────────────────────────────────────────────────────────────────────────

def _add_header(section, company_name: str, doc_type: str):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(p.add_run(f"{company_name}  |  {doc_type}"), size_pt=8, color=GRAY)


def _add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(p)
    for run in p.runs:
        run.font.size      = Pt(8)
        run.font.color.rgb = GRAY
        run.font.name      = "Arial"


# ─────────────────────────────────────────────────────────────────────────────
#  PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(
    doc_type: str,
    department: str,
    company_name: str,
    industry: str,
    region: str,
    sections: List[Dict],
) -> bytes:
    """
    Build a professional .docx document and return raw bytes.

    Args:
        doc_type:     e.g. "Content Strategy Document"
        department:   e.g. "Marketing"
        company_name: e.g. "Turabit Technologies"
        industry:     e.g. "Technology / SaaS"
        region:       e.g. "India"
        sections:     [{"name": "Section Name", "content": "..."}]
                      Content supports: plain text, pipe tables, mermaid blocks,
                      bullet lists, numbered lists.

    Returns:
        bytes: Raw .docx — pass to st.download_button() directly.
    """
    doc = Document()

    # Page setup: US Letter, 1-inch margins
    for section in doc.sections:
        section.page_width    = Twips(12240)
        section.page_height   = Twips(15840)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        _add_header(section, company_name, doc_type)
        _add_footer(section)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # Document title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(8)
    _set_run_font(title_p.add_run(doc_type), size_pt=22, bold=True, color=ACCENT)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(6)
    _set_run_font(
        sub_p.add_run(f"{company_name}  ·  {department}  ·  {industry}  ·  {region}"),
        size_pt=10, color=GRAY
    )

    # Divider
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(16)
    _add_bottom_border(div_p)

    # Sections
    for idx, sec in enumerate(sections):
        name    = sec.get("name", "").strip()
        content = sec.get("content", "").strip()

        if not content:
            continue

        if idx > 0:
            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(6)

        # Section heading with underline
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(5)
        _set_run_font(h.add_run(name), size_pt=13, bold=True, color=ACCENT)
        _add_bottom_border(h)

        # Section body
        _render_section_content(doc, content, section_name=name)

    # End note
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        note_p.add_run(f"{doc_type}  ·  Generated by DocForge AI  ·  Confidential"),
        size_pt=8, color=GRAY
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()